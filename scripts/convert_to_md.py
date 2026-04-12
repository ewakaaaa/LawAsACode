"""
Konwersja dokumentów legislacyjnych (DOCX, PDF, DOC) na format Markdown.

Użycie:
    python scripts/convert_to_md.py data/arrangements/ outputs/arrangements/
"""

import argparse
import os
import subprocess
import sys
from pathlib import Path


def convert_docx_to_md(filepath: str) -> str:
    """Konwertuje plik DOCX na Markdown."""
    import docx

    doc = docx.Document(filepath)
    parts = []

    # Paragraphs and tables in document order
    # python-docx doesn't expose mixed order natively,
    # so we iterate the XML body children
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in body:
        if child.tag == qn("w:p"):
            para = _find_paragraph(doc, child)
            if para is not None:
                text = _convert_paragraph(para)
                if text:
                    parts.append(text)
        elif child.tag == qn("w:tbl"):
            table = _find_table(doc, child)
            if table is not None:
                md_table = _convert_table(table)
                if md_table:
                    parts.append(md_table)

    return "\n\n".join(parts)


def _find_paragraph(doc, element):
    """Znajduje obiekt Paragraph odpowiadający elementowi XML."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    """Znajduje obiekt Table odpowiadający elementowi XML."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _convert_paragraph(para) -> str:
    """Konwertuje paragraf DOCX na Markdown."""
    text = para.text.strip()
    if not text:
        return ""

    style_name = (para.style.name or "").lower() if para.style else ""

    # Heading styles
    if "heading 1" in style_name or "nagłówek 1" in style_name:
        return f"# {text}"
    if "heading 2" in style_name or "nagłówek 2" in style_name:
        return f"## {text}"
    if "heading 3" in style_name or "nagłówek 3" in style_name:
        return f"### {text}"

    # Bold/italic inline formatting - merge consecutive runs with same style
    groups = []  # list of (style, text) where style is "", "b", "i", "bi"
    for run in para.runs:
        t = run.text
        if not t:
            continue
        style = ""
        if run.bold and run.italic:
            style = "bi"
        elif run.bold:
            style = "b"
        elif run.italic:
            style = "i"
        if groups and groups[-1][0] == style:
            groups[-1] = (style, groups[-1][1] + t)
        else:
            groups.append((style, t))

    if groups:
        parts = []
        for style, t in groups:
            if style == "bi":
                parts.append(f"***{t}***")
            elif style == "b":
                parts.append(f"**{t}**")
            elif style == "i":
                parts.append(f"*{t}*")
            else:
                parts.append(t)
        return "".join(parts)
    return text


def _convert_table(table) -> str:
    """Konwertuje tabelę DOCX na Markdown, obsługując scalone komórki."""
    rows_data = []
    for row in table.rows:
        cells = []
        prev_text = None
        for cell in row.cells:
            text = cell.text.strip().replace("\n", " ")
            # Deduplikuj scalone komórki
            if text == prev_text:
                continue
            prev_text = text
            cells.append(text)
        rows_data.append(cells)

    if not rows_data:
        return ""

    return _build_md_table(rows_data)


def convert_pdf_to_md(filepath: str) -> str:
    """Konwertuje plik PDF na Markdown."""
    import fitz

    doc = fitz.open(filepath)
    parts = []

    for page_num, page in enumerate(doc):
        # Extract tables first
        tables = page.find_tables()
        table_rects = []
        table_texts = []

        for table in tables.tables:
            md_table = _convert_pdf_table(table)
            if md_table:
                table_texts.append(md_table)
                table_rects.append(table.bbox)

        # Extract text blocks, skipping areas covered by tables
        blocks = page.get_text("blocks", sort=True)
        page_parts = []

        for block in blocks:
            x0, y0, x1, y1, text, block_no, block_type = block
            if block_type != 0:  # skip images
                continue

            # Skip if block overlaps with a table
            skip = False
            for trect in table_rects:
                if _rects_overlap((x0, y0, x1, y1), trect):
                    skip = True
                    break
            if skip:
                continue

            text = text.strip()
            if text:
                page_parts.append(text)

        # Interleave: first page text, then tables
        if page_parts:
            parts.append("\n\n".join(page_parts))
        for tt in table_texts:
            parts.append(tt)

    doc.close()
    return "\n\n".join(parts)


def _rects_overlap(r1, r2) -> bool:
    """Sprawdza czy dwa prostokąty się nakładają."""
    x0a, y0a, x1a, y1a = r1
    x0b, y0b, x1b, y1b = r2
    return not (x1a < x0b or x1b < x0a or y1a < y0b or y1b < y0a)


def _convert_pdf_table(table) -> str:
    """Konwertuje tabelę z PDF (pymupdf) na Markdown."""
    data = table.extract()
    if not data:
        return ""

    # Clean cells
    cleaned = []
    for row in data:
        cleaned_row = []
        prev = None
        for cell in row:
            text = str(cell).strip().replace("\n", " ") if cell else ""
            # Deduplikuj scalone komórki (powtarzający się tekst obok siebie)
            if text == prev:
                continue
            prev = text
            cleaned_row.append(text)
        cleaned.append(cleaned_row)

    if not cleaned:
        return ""

    return _build_md_table(cleaned)


def _build_md_table(rows_data: list) -> str:
    """Buduje tabelę Markdown z listy wierszy, usuwając puste kolumny."""
    if not rows_data:
        return ""

    # Normalize column count
    max_cols = max(len(r) for r in rows_data)
    for row in rows_data:
        while len(row) < max_cols:
            row.append("")

    # Remove columns that are empty in all rows
    cols_to_keep = []
    for col_idx in range(max_cols):
        if any(row[col_idx].strip() for row in rows_data):
            cols_to_keep.append(col_idx)

    if not cols_to_keep:
        return ""

    rows_data = [[row[i] for i in cols_to_keep] for row in rows_data]
    max_cols = len(cols_to_keep)

    # Skip empty rows
    rows_data = [r for r in rows_data if any(cell.strip() for cell in r)]
    if not rows_data:
        return ""

    # Build markdown table
    lines = []
    header = rows_data[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in rows_data[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def convert_doc_to_md(filepath: str) -> str:
    """Konwertuje plik DOC na Markdown (przez macOS textutil)."""
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", filepath],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        print(f"BŁĄD textutil dla {filepath}: {result.stderr}", file=sys.stderr)
        return ""

    return result.stdout.strip()


def convert_file(filepath: str) -> str:
    """Konwertuje plik na Markdown w zależności od rozszerzenia."""
    ext = Path(filepath).suffix.lower()
    if ext == ".docx":
        return convert_docx_to_md(filepath)
    elif ext == ".pdf":
        return convert_pdf_to_md(filepath)
    elif ext == ".doc":
        return convert_doc_to_md(filepath)
    elif ext == ".xml":
        # Pliki XAdES to podpisy elektroniczne - pomijamy
        return ""
    else:
        print(f"Nieobsługiwany format: {filepath}", file=sys.stderr)
        return ""


def main():
    parser = argparse.ArgumentParser(
        description="Konwersja dokumentów legislacyjnych na Markdown"
    )
    parser.add_argument("input_dir", help="Folder z dokumentami źródłowymi")
    parser.add_argument("output_dir", help="Folder na pliki Markdown")
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    for filepath in sorted(input_dir.iterdir()):
        if filepath.name.startswith("."):
            continue
        if filepath.suffix.lower() == ".xml":
            print(f"Pomijam {filepath.name} (podpis elektroniczny)")
            continue

        print(f"Konwertuję {filepath.name}...")
        md_content = convert_file(str(filepath))

        if md_content:
            output_path = output_dir / (filepath.stem + ".md")
            output_path.write_text(md_content, encoding="utf-8")
            print(f"  → {output_path}")
        else:
            print(f"  ⚠ Pusty wynik dla {filepath.name}")


if __name__ == "__main__":
    main()
